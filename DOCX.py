from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
# 新建空白文档
doc1 = Document()
doc1.add_heading('AAA ',1)
# 创建段落描述

p=doc1.add_paragraph('ASSDDD')
p.add_run('This Font is Times New Roman ').bold=True
a = p.add_run('Setting the front NoW')
a.font.size = Pt(24)
a.font.name ='Times New Roman'
a.font.color.rgb = RGBColor(0, 255, 0)
a.font.highlight_color = WD_COLOR_INDEX.YELLOW
doc1.add_paragraph('HEllo')
#italic,underline is used the same as bold command
doc1.save('word1.docx')




print(type(a))
doc2 = Document('word2.docx')

# 读取每段内容

pl = [ paragraph.text for paragraph in doc2.paragraphs]

for i in pl:
    print(i)
for paragraph in doc2.paragraphs:

    print(paragraph.text)
doc2.save('word2.docx')

